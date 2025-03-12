from datetime import date
import io
import streamlit as st
from docx import Document


from employee import Period, Employee

# Заголовок приложения
st.title("Калькулятор стажа сотрудника")

# Ввод имени сотрудника
employee_name = st.text_input("Введите ФИО сотрудника", placeholder="Иван Иванов")

# Инициализация списка периодов работы в session_state
# session_state используется для сохранения состояния между перезагрузками страницы
if "periods" not in st.session_state:
    st.session_state.periods = []

# Ввод дат периода работы
start_date = st.date_input("Дата трудоустройства", value=None, min_value=date(1970, 1, 1))
end_date = st.date_input("Дата увольнения", value=None, min_value=date(1970, 1, 1))

# Словарь с коэффициентами для льготного стажа
privilegies = {'без льготы': 0, '1 месяц работы за 2 месяца': 2, '1 месяц работы за 1,5 месяца': 1.5}
# Выпадающее окно для выбора коэффициента льготы
privilege = st.selectbox("Выберите коэффициент", privilegies.keys())

# Если имя сотрудника и даты введены, можно добавить период работы
if employee_name and start_date and end_date:
    # Кнопка для добавления периода работы в список
    if st.button("Добавить введенный период работы"):
        # Создаем объект Period и добавляем его в session_state.periods
        st.session_state.periods.append(Period(start_date, end_date, privilegies[privilege]))

# Если список периодов не пуст, отображаем кнопку для расчёта стажа
if st.session_state.periods:
    if st.button("Рассчитать стаж"):
        # Создаем объект Employee и рассчитываем общий стаж
        employee = Employee(employee_name, st.session_state.periods)
        employee.calculate_total_expirience()
        # Выводим результат расчёта стажа
        st.write(employee)

        doc = Document()
        doc.add_heading(employee_name, 0)
        # Сохраняем документ в байтовый объект
        doc.add_heading('Общий стаж сотрудника:', 1)
        doc.add_paragraph(f'Kоличество лет: {employee.year}')
        doc.add_paragraph(f'Kоличество месяцев: {employee.month}')
        doc.add_paragraph(f'Kоличество дней: {employee.day}')
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)

        # Кнопка для скачивания файла
        st.download_button(
            label="Скачать word документ",  # Текст на кнопке
            data=doc_bytes,  # Данные для скачивания
            file_name="doc.docx",  # Имя файла
            mime="text/plain"  # MIME-тип файла (для текстового файла)
        )

# Если список периодов не пуст, отображаем кнопку для очистки списка
if st.session_state.periods:
    if st.button("Очистить список периодов работы"):
        # Очищаем список периодов
        st.session_state.periods = []

# Отображение добавленных периодов работы
st.subheader("Добавленные периоды работы")
for period in st.session_state.periods:
    st.write(period)
